"""
文档解析模块 (DocumentParser)
支持：PDF、DOC、DOCX 格式
功能：自动识别文件格式并提取文本内容
"""

import os
from pathlib import Path
from app.utils.logger import setup_logger

logger = setup_logger(__name__)


class DocumentParser:
    """文档解析器 - 支持多种格式"""
    
    SUPPORTED_FORMATS = {'.pdf', '.doc', '.docx', '.txt'}
    
    @staticmethod
    def parse(file_path: str) -> str:
        """
        自动识别文件格式并解析
        
        Args:
            file_path: 文件路径
        
        Returns:
            提取的文本内容
        
        Raises:
            ValueError: 不支持的文件格式
            FileNotFoundError: 文件不存在
        """
        logger.info(f"[CALL] DocumentParser.parse(file_path={file_path})")
        # 验证文件存在
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 获取文件扩展名
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in DocumentParser.SUPPORTED_FORMATS:
            raise ValueError(f"不支持的文件格式: {file_ext}。支持的格式: {DocumentParser.SUPPORTED_FORMATS}")
        
        # 根据格式调用对应的解析器
        if file_ext == '.pdf':
            return DocumentParser._parse_pdf(file_path)
        elif file_ext == '.docx':
            return DocumentParser._parse_docx(file_path)
        elif file_ext == '.doc':
            return DocumentParser._parse_doc(file_path)
        elif file_ext == '.txt':
            return DocumentParser._parse_txt(file_path)
        else:
            raise ValueError(f"未知的文件格式: {file_ext}")
    
    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """
        解析PDF文档
        
        使用PyPDF2库提取PDF中的文本
        """
        logger.info(f"[CALL] DocumentParser._parse_pdf(file_path={file_path})")
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError(
                "需要安装 PyPDF2 来解析PDF文档。\n"
                "请运行: pip install PyPDF2"
            )
        
        try:
            reader = PdfReader(file_path)
            text = ""
            
            # 逐页提取文本
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    print(f"[警告] 第 {page_num + 1} 页提取文本失败: {e}")
            
            if not text.strip():
                raise ValueError("PDF文档中未找到可提取的文本")
            
            return text.strip()
        
        except Exception as e:
            raise RuntimeError(f"PDF解析失败: {str(e)}")
    
    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """
        解析DOCX文档
        
        使用python-docx库提取DOCX中的文本
        """
        logger.info(f"[CALL] DocumentParser._parse_docx(file_path={file_path})")
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "需要安装 python-docx 来解析DOCX文档。\n"
                "请运行: pip install python-docx"
            )
        
        try:
            doc = Document(file_path)
            text = ""
            
            # 提取所有段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # 提取表格中的文本（可选）
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise ValueError("DOCX文档中未找到可提取的文本")
            
            return text.strip()
        
        except Exception as e:
            raise RuntimeError(f"DOCX解析失败: {str(e)}")
    
    @staticmethod
    def _parse_doc(file_path: str) -> str:
        """
        解析DOC文档
        
        使用python-pptx或其他库来处理DOC格式
        注：DOC格式解析较复杂，这里提供基础实现
        """
        logger.info(f"[CALL] DocumentParser._parse_doc(file_path={file_path})")
        try:
            # 首先尝试使用 python-doc（如果安装）
            try:
                import docx2docx
                return DocumentParser._parse_doc_with_docx2docx(file_path)
            except ImportError:
                pass
            
            # 备选方案：使用 python-docx（通常能处理较新的DOC文件）
            from docx import Document
            doc = Document(file_path)
            text = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            if not text.strip():
                raise ValueError("DOC文档中未找到可提取的文本")
            
            return text.strip()
        
        except ImportError:
            raise ImportError(
                "需要安装库来解析DOC文档。\n"
                "推荐: pip install python-docx\n"
                "或: pip install docx2docx"
            )
        except Exception as e:
            raise RuntimeError(f"DOC解析失败: {str(e)}")
    
    @staticmethod
    def _parse_doc_with_docx2docx(file_path: str) -> str:
        """使用 docx2docx 库解析 DOC"""
        logger.info(f"[CALL] DocumentParser._parse_doc_with_docx2docx(file_path={file_path})")
        try:
            import docx2docx
            # 将DOC转换为DOCX，再用python-docx处理
            from docx import Document
            
            # 临时转换
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp_path = tmp.name
            
            # 转换
            docx2docx.convert(file_path, tmp_path)
            
            # 解析DOCX
            doc = Document(tmp_path)
            text = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # 清理临时文件
            os.unlink(tmp_path)
            
            return text.strip()
        except Exception as e:
            raise RuntimeError(f"使用docx2docx解析DOC失败: {str(e)}")
    
    @staticmethod
    def _parse_txt(file_path: str) -> str:
        """
        解析TXT文本文件
        
        支持UTF-8和GBK编码
        """
        logger.info(f"[CALL] DocumentParser._parse_txt(file_path={file_path})")
        try:
            # 首先尝试UTF-8
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text.strip()
        except UnicodeDecodeError:
            # 如果UTF-8失败，尝试GBK
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    text = f.read()
                return text.strip()
            except Exception as e:
                raise RuntimeError(f"TXT解析失败（尝试UTF-8和GBK均失败）: {str(e)}")
        except Exception as e:
            raise RuntimeError(f"TXT解析失败: {str(e)}")
    
    @staticmethod
    def get_file_format(file_path: str) -> str:
        """
        获取文件格式
        
        Returns:
            文件格式字符串（'pdf', 'doc', 'docx', 'txt'）
        """
        logger.info(f"[CALL] DocumentParser.get_file_format(file_path={file_path})")
        ext = Path(file_path).suffix.lower().lstrip('.')
        return ext if ext in DocumentParser.SUPPORTED_FORMATS else 'unknown'
